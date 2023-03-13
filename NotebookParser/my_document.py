from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import paragraphs


document = Document()
for section in document.sections:
    section.left_margin = Inches(0.98)
    section.right_margin = Inches(0.59)
    section.top_margin = Inches(0.49)
    section.bottom_margin = Inches(0.59)

paragraph_format = document.styles['Normal'].paragraph_format
paragraph_format.line_spacing = 1.5
paragraph_format.first_line_indent = Inches(0.492)


para = document.add_paragraph()
para.text = 'dssssssssssssssssssssssfffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffff'
paragraph_format.space_after

para = document.add_paragraph()
picture_run = para.add_run()
picture_run.add_picture('index.png')
para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
para.add_run('\nImage Text')
# para = document.add_heading('3 Виконання', 1)
# para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


document.save('test.docx')